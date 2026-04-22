
import os
import mimetypes
from typing import List, Optional, Tuple
from pathlib import Path
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from docx import Document as DocxDocument


class DocumentProcessor:
    """Process multi-modal documents (PDF, images, text) for RAG ingestion."""
    
    def __init__(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        use_ocr: bool = True,
        tesseract_path: Optional[str] = None
    ):
        """
        Initialize document processor with multi-modal support.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            use_ocr: Enable OCR for images and scanned PDFs
            tesseract_path: Path to Tesseract executable (if not in PATH)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_ocr = use_ocr
        
        # Configure Tesseract if custom path provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Advanced text splitter with semantic separators
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=[
                "\n\n\n",  # Multiple line breaks (major sections)
                "\n\n",    # Paragraphs
                "\n",      # Lines
                ". ",      # Sentences
                ", ",      # Clauses
                " ",       # Words
                ""         # Characters
            ]
        )
        
        mimetypes.init()
    
    def detect_file_type(self, file_path: str) -> str:
        """
        Detect file type using mimetypes and extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type category: 'pdf', 'image', 'text', 'docx', 'unknown'
        """
        mime_type, _ = mimetypes.guess_type(file_path)
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf' or (mime_type and 'pdf' in mime_type):
            return 'pdf'
        elif extension in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.gif'] or (mime_type and 'image' in mime_type):
            return 'image'
        elif extension in ['.txt', '.md', '.csv'] or (mime_type and 'text' in mime_type):
            return 'text'
        elif extension == '.docx':
            return 'docx'
        else:
            return 'unknown'
    
    def extract_text_from_pdf(self, pdf_path: str) -> Tuple[str, bool]:
        """
        Extract text from a PDF file with OCR fallback.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Tuple of (extracted text, used_ocr)
        """
        try:
            reader = PdfReader(pdf_path)
            text = ""
            
            for page in reader.pages:
                page_text = page.extract_text()
                text += page_text + "\n"
            
            # Check if PDF appears to be scanned (very little text extracted)
            if self.use_ocr and len(text.strip()) < 100:
                print(f"  PDF appears scanned, using OCR...")
                return self.extract_text_with_ocr_from_pdf(pdf_path), True
            
            return text, False
        except Exception as e:
            raise Exception(f"Error reading PDF {pdf_path}: {str(e)}")
    
    def extract_text_with_ocr_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF using OCR.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text as a string
        """
        try:
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=300)
            text = ""
            
            for i, image in enumerate(images):
                print(f"    Processing page {i+1}/{len(images)} with OCR...")
                page_text = pytesseract.image_to_string(image)
                text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error performing OCR on PDF {pdf_path}: {str(e)}")
    
    def extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from an image using OCR.
        
        Args:
            image_path: Path to the image file
            
        Returns:
            Extracted text as a string
        """
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            return text
        except Exception as e:
            raise Exception(f"Error reading image {image_path}: {str(e)}")
    
    def extract_text_from_text_file(self, text_path: str) -> str:
        """
        Extract text from a plain text file.
        
        Args:
            text_path: Path to the text file
            
        Returns:
            File contents as a string
        """
        try:
            with open(text_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Error reading text file {text_path}: {str(e)}")
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text as a string
        """
        try:
            doc = DocxDocument(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            raise Exception(f"Error reading DOCX file {docx_path}: {str(e)}")
    
    def process_document(self, file_path: str) -> List[Document]:
        """
        Process any supported document type into chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of Document objects with chunks
        """
        file_type = self.detect_file_type(file_path)
        filename = os.path.basename(file_path)
        
        # Extract text based on file type
        if file_type == 'pdf':
            text, used_ocr = self.extract_text_from_pdf(file_path)
            doc_type = 'pdf_ocr' if used_ocr else 'pdf'
        elif file_type == 'image':
            if not self.use_ocr:
                raise ValueError(f"OCR is disabled. Cannot process image: {filename}")
            text = self.extract_text_from_image(file_path)
            doc_type = 'image_ocr'
        elif file_type == 'text':
            text = self.extract_text_from_text_file(file_path)
            doc_type = 'text'
        elif file_type == 'docx':
            text = self.extract_text_from_docx(file_path)
            doc_type = 'docx'
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Validate extracted text
        if not text or len(text.strip()) < 10:
            raise ValueError(f"No meaningful text extracted from {filename}")
        
        # Create initial document with rich metadata
        doc = Document(
            page_content=text,
            metadata={
                "source": filename,
                "file_path": file_path,
                "doc_type": doc_type,
                "file_size": os.path.getsize(file_path),
                "char_count": len(text)
            }
        )
        
        # Split into chunks with semantic boundaries
        chunks = self.text_splitter.split_documents([doc])
        
        # Enhance chunk metadata and normalize content for better case-insensitive retrieval
        for i, chunk in enumerate(chunks):
            # Store both original and normalized versions for flexible retrieval
            # The vector embeddings will work better with normalized text
            normalized_content = chunk.page_content.lower()
            chunk.page_content = normalized_content
            
            chunk.metadata.update({
                "chunk_id": i,
                "chunk_total": len(chunks),
                "chunk_chars": len(chunk.page_content)
            })
        
        return chunks
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[Document]:
        """
        Process multiple document files of various types.
        
        Args:
            file_paths: List of paths to document files
            
        Returns:
            List of all document chunks with metadata
        """
        all_chunks = []
        
        for file_path in file_paths:
            print(f"Processing: {os.path.basename(file_path)}")
            try:
                chunks = self.process_document(file_path)
                all_chunks.extend(chunks)
                print(f"  ✓ Created {len(chunks)} chunks from {chunks[0].metadata['doc_type']}")
            except Exception as e:
                print(f"  ✗ Error: {str(e)}")
                continue
        
        return all_chunks
